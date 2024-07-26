import os
import random
from flask import Flask, render_template, request
from docx import Document
from werkzeug.utils import secure_filename

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/products')
def products():
    return render_template('index.html')

@app.route('/products/<pid>')
def product(pid):
    return "product: %s" % pid

@app.route('/products/<pid>/downloadfile/<filename>')
def downloadFile(pid, filename):
    return "product: %s - filename %s" % pid, filename

@app.route('/productss/xss/<name>')
def productXSS(name):
    return "%s" % name



@app.route('/login')
def login():
    return 'login'

@app.route('/user')
def userHome():
    return 'userHome'

@app.route('/admin')
def adminHome():
    return 'adminHome'

#
#
#


@app.route("/xxe")
def xxe():
    return render_template('test.html')

@app.route("/xxe_uploader", methods=['GET', 'POST'])  # /<string:name>/")
def xxe_uploader():
    if request.method == 'POST':

        f = request.files['file']
        rand = random.randint(1, 100)
        fname = secure_filename(f.filename)
        fname = str(rand) + fname  # change file name
        cwd = os.getcwd()
        file_path = cwd + '/Files/' + fname
        f.save(file_path)  # save file locally

        # Access saved file
        document = Document(file_path)
        for para in document.paragraphs:
            print (para.text)  # '\n\n'.join([para.text for paragraph in document.paragraphs])

    # return "file uploaded successfully"
    return render_template('view.html', name=para.text)



if __name__ == '__main__':
    app.jinja_env.auto_reload = True
    app.config['TEMPLATES_AUTO_RELOAD'] = True
    app.run(debug=True, host='0.0.0.0')